import os
from io import StringIO, BytesIO
import datetime

#django module
from django.forms.forms import Form
from django.views.generic import CreateView, ListView, UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect
from django.contrib import messages
from django.shortcuts import redirect
from django.conf import settings

#3rd party module
from docx import Document

from .models import Corona3
from .forms import BasicDataForm, minDataForm


class BasicFormAddNewView(CreateView):

    model = Corona3
    form_class = BasicDataForm
    template_name = 'Corona3/basicData.html' 
    success_url = '/CRN3/BasicData/'
    
    def form_invalid(self, form):
        print(self.request, form.errors)
        return super().form_invalid(form)
    
    def post(self, request):
        super(BasicFormAddNewView, self).post(request)
        messages.success(request,"บันทึกข้อมูลเรียบร้อย")
        return redirect('/CRN3/BasicData/')

class minDataAddNewView(CreateView):

    model = Corona3
    form_class = minDataForm
    template_name = 'Corona3/minData.html' 
    success_url = '/CRN3/minData/'
    
    def form_invalid(self, form):
        print(self.request, form.errors)
        return super().form_invalid(form)
    
    def post(self, request):
        super(minDataAddNewView, self).post(request)
        messages.success(request,"บันทึกข้อมูลเรียบร้อยแล้ว")
        return redirect('/CRN3/minData/')

class Corona3ListView(LoginRequiredMixin, ListView):
    login_url = '/login'
    model = Corona3
    template_name = 'Corona3/list.html'
    ordering = ['-DateReport','-id']    
    paginate_by = 10

class Corona3UpdateView(LoginRequiredMixin,UpdateView):
    model = Corona3
    form_class = BasicDataForm
    template_name = 'Corona3/basicData.html'
    success_url = '/CRN3/list/'


def corona3_Document(request,pk):
    # testdoc = static('documents/form_corona3.docx')
    testdoc =  os.path.join(settings.TEMPLATES[0]['DIRS'][0],'documents/form_corona3.docx')
    # testdoc =  "/home/sammy/afcovid/templates/documents/form_corona.docx"

    document = Document(testdoc)

    corona3 = Corona3.objects.get(id = pk)
    docx_title= f"Corona3-{pk}.docx"

    time_and_date = "2019-10-21"
    employee_name = corona3.FullName
    manager_name = "Michelle Johnson"

    dic = {
            '#Fullname#':corona3.FullName,
            '#PersonID#':corona3.PersonID,
            '#Gender#':corona3.get_Gender_display(),
<<<<<<< HEAD
            '#Age#':corona3.Age,
=======
            '#Age#':str(corona3.Age),
>>>>>>> DevMuek
            '#Nationority#':corona3.get_Nationality_display(),
            '#CareerPatient#':corona3.CareerPatient,
            '#Mobile#':corona3.Mobile,
            '#PlaceWork#':corona3.PlaceWork,
            '#Address#':corona3.Address,
<<<<<<< HEAD
            '#Swine#':corona3.Swine,
            '#SubDistrict#':corona3.SubDistrict,
            '#District#':corona3.District,
            '#Province#':corona3.get_Province_display(),  
            '#DatePatient#':corona3.DatePatient,
            '#DateFirstTreat#':corona3.DateFirstTreat,
            '#DateDiagnose#':corona3.DateDiagnose,
            '#CaseFemale#':corona3.CaseFemale,
            '#DateCheckRTPCR#':corona3.DateCheckRTPCR,
            '#TypeExampleRTPCR#':corona3.TypeExampleRTPCR,
            '#PlaceCheckRTCR#':corona3.PlaceCheckRTPCR,
            '#ResultsCheckRTPCR#':corona3.ResultsCheckRTPCR,
            '#DateCheckAntigen#':corona3.DateCheckAntigen,
=======
            '#Swine#':str(corona3.Swine),
            '#SubDistrict#':corona3.SubDistrict,
            '#District#':corona3.District,
            '#Province#':corona3.get_Province_display(),  
            '#DatePatient#':str(corona3.DatePatient),
            '#DateFirstTreat#':str(corona3.DateFirstTreat),
            '#DateDiagnose#':str(corona3.DateDiagnose),
            '#CaseFemale#':str(corona3.CaseFemale),
            '#DateCheckRTPCR#':str(corona3.DateCheckRTPCR),
            '#TypeExampleRTPCR#':str(corona3.TypeExampleRTPCR),
            '#PlaceCheckRTCR#':corona3.PlaceCheckRTPCR,
            '#ResultsCheckRTPCR#':corona3.ResultsCheckRTPCR,
            '#DateCheckAntigen#':str(corona3.DateCheckAntigen),
>>>>>>> DevMuek
            '#TypeExampleAntigen#':corona3.TypeExampleAntigen,
            '#PlaceCheckAntigen#':corona3.PlaceCheckAntigen,
            '#ResultsCheckAntigen#':corona3.ResultsCheckAntigen,
            '#HospitalProvince#':corona3.get_HospitalProvince_display(),
            '#NameHospitalTreat#':corona3.NameHospitalTreat,
<<<<<<< HEAD
            '#DateCheckAntibody1#':corona3.DateCheckAntibody1,
            '#PlaceCheckAntibody1#':corona3.PlaceCheckAntibody1,
            '#TypeExampleAntibody1#':corona3.TypeExampleAntibody1,
            '#DateCheckAntibody2#':corona3.DateCheckAntibody2,
            '#TypeExampleAntibody2#':corona3.TypeExampleAntibody2,
            '#PlaceCheckAntibody2#':corona3.PlaceCheckAntibody2,
            '#ReceivedVaccine#':corona3.ReceivedVaccine,
            '#BookReceivedVaccine#':corona3.BookReceivedVaccine,
            '#DateReceivedVaccine1#':corona3.DateReceivedVaccine1,
            '#NameVaccine1#':corona3.NameVaccine1,
            '#PlaceReceivedVaccine1#':corona3.PlaceReceivedVaccine1,
            '#DateReceivedVaccine2#':corona3.DateReceivedVaccine2,
            '#NameVaccine2#':corona3.NameVaccine2,
            '#PlaceReceivedVaccine2#':corona3.PlaceReceivedVaccine2,
            '#LiveInCovid#':corona3.LiveInCovid,
            '#InThaiProvince#':corona3.InThaiProvince,
            '#InForeignCountry#':corona3.InForeignCountry,
            '#InForeignCity#':corona3.InForeignCity,
            '#NearCovid#':corona3.NearCovid,
            '#ContactCovid#':corona3.ContactCovid,
            '#ContactCovidText#':corona3.ContactCovidText,
            '#CareerNearCovid#':corona3.CareerNearCovid,
            '#TravelInCovid#':corona3.TravelInCovid,
            '#TravelInCovidText#':corona3.TravelInCovidText,
            '#AuthoritiesMedical#':corona3.AuthoritiesMedical,
            '#HistoryRisky#':corona3.HistoryRisky,
            '#DiseasePatient#':corona3.DiseasePatient,
            '#HistoryRiskyText#':corona3.HistoryRiskyText,
            '#ContactRisky#':corona3.ContactRisky,
            '#ContactRiskyTrace#':corona3.ContactRiskyTrace,
            '#PlaceConfineContactRisky1#':corona3.PlaceConfineContactRisky1,
            '#PlaceConfineContactRisky2#':corona3.PlaceConfineContactRisky2,
            '#ContactLowRisk#':corona3.ContactLowRisk,
            '#ContactLowRiskTrace#':corona3.ContactLowRiskTrace,
            '#PlaceConfineContactLowRisk1#':corona3.PlaceConfineContactLowRisk1,
            '#PlaceConfineContactLowRisk2#':corona3.PlaceConfineContactLowRisk2,
            '#UserReport#':corona3.UserReport,
            '#Unit#':corona3.Unit,
            '#Mobile2#':corona3.Mobile2,
            '#DateReport#':corona3.DateReport
=======
            '#DateCheckAntibody1#':str(corona3.DateCheckAntibody1),
            '#PlaceCheckAntibody1#':corona3.PlaceCheckAntibody1,
            '#TypeExampleAntibody1#':corona3.TypeExampleAntibody1,
            '#DateCheckAntibody2#':str(corona3.DateCheckAntibody2),
            '#TypeExampleAntibody2#':corona3.TypeExampleAntibody2,
            '#PlaceCheckAntibody2#':corona3.PlaceCheckAntibody2,
            '#ReceivedVaccine#':str(corona3.ReceivedVaccine),
            '#BookReceivedVaccine#':str(corona3.BookReceivedVaccine),
            '#DateReceivedVaccine1#':str(corona3.DateReceivedVaccine1),
            '#NameVaccine1#':str(corona3.NameVaccine1),
            '#PlaceReceivedVaccine1#':corona3.PlaceReceivedVaccine1,
            '#DateReceivedVaccine2#':str(corona3.DateReceivedVaccine2),
            '#NameVaccine2#':str(corona3.NameVaccine2),
            '#PlaceReceivedVaccine2#':corona3.PlaceReceivedVaccine2,
            '#LiveInCovid#':str(corona3.LiveInCovid),
            '#InThaiProvince#':str(corona3.InThaiProvince),
            '#InForeignCountry#':corona3.InForeignCountry,
            '#InForeignCity#':corona3.InForeignCity,
            '#NearCovid#':str(corona3.NearCovid),
            '#ContactCovid#':str(corona3.ContactCovid),
            '#ContactCovidText#':str(corona3.ContactCovidText),
            '#CareerNearCovid#':str(corona3.CareerNearCovid),
            '#TravelInCovid#':str(corona3.TravelInCovid),
            '#TravelInCovidText#':str(corona3.TravelInCovidText),
            '#AuthoritiesMedical#':str(corona3.AuthoritiesMedical),
            '#HistoryRisky#':str(corona3.HistoryRisky),
            '#DiseasePatient#':str(corona3.DiseasePatient),
            '#HistoryRiskyText#':str(corona3.HistoryRiskyText),
            '#ContactRisky#':str(corona3.ContactRisky),
            '#ContactRiskyTrace#':str(corona3.ContactRiskyTrace),
            '#PlaceConfineContactRisky1#':str(corona3.PlaceConfineContactRisky1),
            '#PlaceConfineContactRisky2#':str(corona3.PlaceConfineContactRisky2),
            '#ContactLowRisk#':str(corona3.ContactLowRisk),
            '#ContactLowRiskTrace#':str(corona3.ContactLowRiskTrace),
            '#PlaceConfineContactLowRisk1#':str(corona3.PlaceConfineContactLowRisk1),
            '#PlaceConfineContactLowRisk2#':str(corona3.PlaceConfineContactLowRisk2),
            '#UserReport#':corona3.UserReport,
            '#Unit#':corona3.Unit,
            '#Mobile2#':corona3.Mobile2,
            '#DateReport#':str(corona3.DateReport),
>>>>>>> DevMuek
            }
    # print(dic)


    for para in document.paragraphs:
        for key, value in dic.items():
            replace_value = value if value != None else ' '
            replace_value = replace_value.strftime("%d %b %Y") if isinstance(replace_value, datetime.date) else replace_value
            replace_value = replace_value if type(replace_value) == str else str(replace_value)
            if key in para.text:
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, replace_value)
                        inline[i].text = text

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response








    


    
